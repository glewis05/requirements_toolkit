# parsers/word_parser.py
# ============================================================================
# PURPOSE: Parse Microsoft Word documents (.docx) containing requirements
#
# Think of this like reading a Word document in R with officer::read_docx(),
# but with specialized logic to extract requirements from various formats:
# - Tables with structured data (ID, Description, Priority columns)
# - Narrative prose with embedded requirements ("The system shall...")
# - Numbered/bulleted lists of requirements
# - Mixed formats within the same document
#
# Word documents are notoriously messy for requirements because different
# authors use different structures. This module handles that chaos.
# ============================================================================

import re
from pathlib import Path
from typing import Optional
from collections import defaultdict

# python-docx is the go-to library for reading Word documents in Python
# R EQUIVALENT: Similar to officer::read_docx() but with more control
# Install with: pip install python-docx
try:
    from docx import Document
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None


class WordParser:
    """
    PURPOSE:
        Parse Microsoft Word documents containing requirements and extract
        them into a normalized list of dictionaries. Handles both structured
        tables and narrative prose formats.

    R EQUIVALENT:
        Similar to using officer::read_docx() + officer::docx_summary() +
        custom parsing logic. In R you might loop through body elements
        and extract text from paragraphs and tables separately.

    WHY A CLASS:
        - Holds document state and parsing context (current section, etc.)
        - Tracks statistics across the parsing process
        - Encapsulates all Word-specific logic in one place

    USAGE:
        parser = WordParser("inputs/word/prd.docx")
        requirements = parser.parse()
        # requirements is a list of dicts matching ExcelParser output

    SUPPORTED FORMATS:
        1. TABLE FORMAT: Requirements in structured tables
           | ID        | Description                | Priority |
           |-----------|----------------------------|----------|
           | REQ-001   | User can login with email  | High     |

        2. PROSE FORMAT: Requirements embedded in paragraphs
           "The system shall allow users to login with email..."
           "Users must be able to reset their password..."

        3. LIST FORMAT: Requirements in numbered/bulleted lists
           1. User registration functionality
           2. Password reset via email
           - Admin can manage users
           - System sends notifications

        4. MIXED FORMAT: Combination of all the above
    """

    def __init__(self, file_path: str) -> None:
        """
        PURPOSE:
            Initialize the parser with a path to a Word document.
            Validates the file exists and has .docx extension.

        R EQUIVALENT:
            Like setting up parameters before calling officer::read_docx()
            path <- "my_file.docx"  # just storing the path

        PARAMETERS:
            file_path (str): Path to the Word document
                            Example: "inputs/word/requirements.docx"

        RETURNS:
            None (constructor)

        RAISES:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a .docx file

        WHY THIS APPROACH:
            We validate early but don't parse yet — allows user to
            configure options before the expensive parsing operation.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for Word parsing. "
                "Install with: pip install python-docx"
            )

        self.file_path = Path(file_path)

        # Validate file exists
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {self.file_path}")

        # Validate extension
        if self.file_path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError(
                f"Invalid file type: {self.file_path.suffix}. "
                "WordParser only supports .docx files."
            )

        # Note: .doc (old Word format) requires different handling
        if self.file_path.suffix.lower() == '.doc':
            raise ValueError(
                "Old .doc format not supported. Please save as .docx first."
            )

        # Will hold the python-docx Document object after loading
        self.document = None

        # Current section context (from headings)
        # WHY: Headings give us category information for requirements
        self._current_section = ""
        self._section_stack = []  # For nested headings

        # Statistics tracking
        # WHY: Helps user understand what was extracted and spot issues
        self.stats = {
            'paragraphs_parsed': 0,
            'tables_parsed': 0,
            'requirements_from_tables': 0,
            'requirements_from_prose': 0,
            'requirements_from_lists': 0,
            'sections_found': 0,
            'comments_extracted': 0,
            'total_requirements': 0
        }

        # ====================================================================
        # CONFIGURATION: Requirement Detection
        # ====================================================================

        # Keywords that indicate a sentence is a requirement
        # WHY: These are standard requirement language patterns
        self.requirement_indicators = [
            'shall',      # "The system shall..."
            'must',       # "Users must be able to..."
            'should',     # "The application should..."
            'will',       # "The feature will..."
            'require',    # "Required: ..." or "This requirement..."
            'as a',       # User story: "As a user, I want..."
            'i want',     # User story continuation
            'so that',    # User story benefit
        ]

        # Patterns that indicate a requirement ID
        # WHY: Many docs have explicit IDs like REQ-001, FR-1.2, etc.
        self.id_patterns = [
            r'^(REQ[-_]?\d+[-.]?\d*)',           # REQ-001, REQ001, REQ-1.2
            r'^(FR[-_]?\d+[-.]?\d*)',            # FR-001 (Functional Req)
            r'^(NFR[-_]?\d+[-.]?\d*)',           # NFR-001 (Non-Functional)
            r'^(UC[-_]?\d+[-.]?\d*)',            # UC-001 (Use Case)
            r'^(US[-_]?\d+[-.]?\d*)',            # US-001 (User Story)
            r'^(BR[-_]?\d+[-.]?\d*)',            # BR-001 (Business Rule)
            r'^([A-Z]{2,5}[-_]\d+[-.]?\d*)',     # Generic: XX-NNN
            r'^(\d+\.\d+(?:\.\d+)?)',            # 1.2.3 (numbered)
        ]

        # Common table header names for requirement columns
        # WHY: Tables use various column names for the same concept
        self.table_column_mappings = {
            'id': ['id', 'req id', 'requirement id', 'ref', 'reference',
                   'number', '#', 'no', 'no.', 'item'],
            'description': ['description', 'requirement', 'requirement text',
                           'text', 'detail', 'details', 'specification',
                           'statement', 'content', 'body'],
            'priority': ['priority', 'pri', 'importance', 'criticality',
                        'urgency', 'moscow', 'must/should/could'],
            'category': ['category', 'type', 'area', 'module', 'component',
                        'section', 'group', 'classification'],
            'status': ['status', 'state', 'progress', 'complete'],
            'notes': ['notes', 'comments', 'remarks', 'additional info',
                     'rationale', 'justification']
        }

        # Priority keyword mappings (same as other parsers for consistency)
        self.priority_keywords = {
            'critical': ['critical', 'must have', 'must-have', 'essential',
                        'mandatory', 'required', 'p1', 'p0', 'urgent',
                        'blocking', 'showstopper'],
            'high': ['high', 'should have', 'should-have', 'important',
                    'p2', 'major', 'significant'],
            'medium': ['medium', 'could have', 'could-have', 'normal',
                      'p3', 'moderate', 'standard'],
            'low': ['low', 'nice to have', 'nice-to-have', 'optional',
                   'p4', 'minor', 'defer', 'future']
        }

    # ========================================================================
    # MAIN PARSING ENTRY POINT
    # ========================================================================

    def parse(self) -> list[dict]:
        """
        PURPOSE:
            Parse the Word document and extract all requirements.
            Handles tables, paragraphs, and lists in the document.

        R EQUIVALENT:
            In R with officer, you'd do something like:
            doc <- read_docx("file.docx")
            content <- docx_summary(doc)
            # Then filter/process content rows...

        RETURNS:
            list[dict]: List of requirement dictionaries, each containing:
                - raw_text: Original text from the document
                - row_number: Paragraph/row number for reference
                - source_cell: Location like "Table 2, Row 3" or "Para 15"
                - requirement_id: Extracted or generated ID
                - priority: Inferred priority (Critical/High/Medium/Low)
                - category: From section headings
                - description: Cleaned requirement text
                - notes: Additional context (comments, etc.)

        WHY THIS APPROACH:
            We process the document in order, maintaining section context
            from headings. This preserves the document's organizational
            structure and gives us category information for free.
        """
        # Load the document
        self.document = Document(str(self.file_path))

        # Storage for extracted requirements
        requirements = []

        # Track paragraph number for reference
        para_number = 0

        # Track table number for reference
        table_number = 0

        # ====================================================================
        # PROCESS DOCUMENT BODY
        # ====================================================================
        # Word documents have a body containing paragraphs and tables
        # We iterate through in document order to maintain context

        # Get all body elements (paragraphs and tables)
        # WHY: python-docx provides separate iterators for each, but we need
        # to process them in document order for section context
        body_elements = self._get_body_elements()

        for element in body_elements:
            if isinstance(element, Paragraph):
                para_number += 1
                self.stats['paragraphs_parsed'] += 1

                # Check if this is a heading (sets section context)
                if self._is_heading(element):
                    self._update_section_context(element)
                    continue

                # Check if this paragraph contains requirements
                para_requirements = self._parse_paragraph(element, para_number)
                requirements.extend(para_requirements)

            elif isinstance(element, Table):
                table_number += 1
                self.stats['tables_parsed'] += 1

                # Parse requirements from the table
                table_requirements = self._parse_table(element, table_number)
                requirements.extend(table_requirements)

        # ====================================================================
        # POST-PROCESSING
        # ====================================================================

        # Extract any comments from the document
        comment_requirements = self._extract_comments()
        requirements.extend(comment_requirements)

        # Deduplicate requirements (sometimes same text appears multiple places)
        requirements = self._deduplicate_requirements(requirements)

        # Update stats
        self.stats['total_requirements'] = len(requirements)

        return requirements

    # ========================================================================
    # DOCUMENT STRUCTURE HELPERS
    # ========================================================================

    def _get_body_elements(self) -> list:
        """
        PURPOSE:
            Get all body elements (paragraphs and tables) in document order.

        R EQUIVALENT:
            Similar to officer::docx_summary() which returns all elements
            in a data frame with their positions.

        RETURNS:
            list: Mixed list of Paragraph and Table objects

        WHY THIS APPROACH:
            python-docx's Document.paragraphs and Document.tables give
            us elements separately. We need to interleave them in the
            order they appear in the document to track section context.
        """
        # Access the document body XML directly
        # WHY: This gives us elements in document order
        body = self.document.element.body
        elements = []

        # Iterate through all child elements of the body
        for child in body:
            # Check if it's a paragraph
            if child.tag == qn('w:p'):
                # Find the corresponding Paragraph object
                for para in self.document.paragraphs:
                    if para._element is child:
                        elements.append(para)
                        break

            # Check if it's a table
            elif child.tag == qn('w:tbl'):
                # Find the corresponding Table object
                for table in self.document.tables:
                    if table._tbl is child:
                        elements.append(table)
                        break

        return elements

    def _is_heading(self, paragraph: Paragraph) -> bool:
        """
        PURPOSE:
            Check if a paragraph is a heading (Heading 1, Heading 2, Title, etc.)

        R EQUIVALENT:
            In officer, you'd check content_type == "heading"
            or check the style name.

        PARAMETERS:
            paragraph (Paragraph): The paragraph to check

        RETURNS:
            bool: True if this is a heading paragraph

        WHY THIS APPROACH:
            Headings in Word are just paragraphs with special styles.
            We check the style name to identify them. This includes
            both "Heading X" styles and "Title"/"Subtitle" styles.
        """
        style_name = paragraph.style.name.lower() if paragraph.style else ''
        heading_styles = ['heading', 'title', 'subtitle', 'toc heading']
        return any(style_name.startswith(hs) for hs in heading_styles)

    def _update_section_context(self, heading: Paragraph) -> None:
        """
        PURPOSE:
            Update the current section context based on a heading.
            This affects the category assigned to subsequent requirements.

        R EQUIVALENT:
            Like maintaining a state variable while iterating through
            a document: current_section <- heading_text

        PARAMETERS:
            heading (Paragraph): The heading paragraph

        WHY THIS APPROACH:
            Headings provide natural categorization for requirements.
            "3.1 Authentication Requirements" tells us the next
            requirements belong to "Authentication" category.
        """
        heading_text = heading.text.strip()

        # Skip empty headings
        if not heading_text:
            return

        # Get heading level from style (Heading 1, Heading 2, etc.)
        style_name = heading.style.name.lower() if heading.style else ''

        # Title and Subtitle are level 0 (document level, not section)
        if style_name.startswith('title') or style_name.startswith('subtitle'):
            # Don't include document title in section context
            self.stats['sections_found'] += 1
            return

        level_match = re.search(r'(\d+)', style_name)
        level = int(level_match.group(1)) if level_match else 1

        # Clean the heading text (remove numbering like "3.1 ")
        clean_heading = re.sub(r'^[\d.]+\s*', '', heading_text)
        clean_heading = clean_heading.strip()

        # Update section stack based on level
        # WHY: Allows us to track nested sections like
        # "3. Requirements" > "3.1 Functional" > "3.1.1 Login"
        while len(self._section_stack) >= level:
            self._section_stack.pop()

        self._section_stack.append(clean_heading)
        self._current_section = ' > '.join(self._section_stack)

        self.stats['sections_found'] += 1

    # ========================================================================
    # PARAGRAPH PARSING
    # ========================================================================

    def _parse_paragraph(self, paragraph: Paragraph, para_number: int) -> list[dict]:
        """
        PURPOSE:
            Extract requirements from a single paragraph.
            Handles prose text, numbered lists, and bulleted lists.

        R EQUIVALENT:
            Like str_extract_all() with regex patterns to find
            requirement sentences within a text block.

        PARAMETERS:
            paragraph (Paragraph): The paragraph to parse
            para_number (int): Paragraph number for reference

        RETURNS:
            list[dict]: List of requirement dictionaries (may be empty)

        WHY THIS APPROACH:
            We check if the paragraph contains requirement indicators,
            then extract the full sentence or list item as the requirement.
        """
        requirements = []
        text = paragraph.text.strip()

        # Skip empty paragraphs
        if not text:
            return requirements

        # Skip very short text (unlikely to be a requirement)
        if len(text) < 10:
            return requirements

        # Check if this is a list item (numbered or bulleted)
        is_list_item = self._is_list_item(paragraph)

        # ====================================================================
        # EXTRACT REQUIREMENT ID IF PRESENT
        # ====================================================================
        explicit_id = self._extract_requirement_id(text)

        # ====================================================================
        # CHECK FOR REQUIREMENT INDICATORS
        # ====================================================================
        text_lower = text.lower()
        has_indicator = any(ind in text_lower for ind in self.requirement_indicators)

        # Filter out non-requirement text (overviews, introductions, etc.)
        if self._is_non_requirement_text(text):
            return requirements

        # If this looks like a requirement or is an explicitly numbered item
        if has_indicator or explicit_id or is_list_item:
            # Parse out any embedded priority
            priority = self._infer_priority(text)

            # Clean the requirement text
            clean_text = self._clean_requirement_text(text)

            # Build the requirement dictionary
            requirement = {
                'raw_text': text,
                'description': clean_text,
                'row_number': para_number,
                'source_cell': f"Paragraph {para_number}",
                'requirement_id': explicit_id or self._generate_id(para_number),
                'priority': priority,
                'category': self._current_section or 'General',
                'notes': '',
                'title': '',
                'status': ''
            }

            requirements.append(requirement)

            if is_list_item:
                self.stats['requirements_from_lists'] += 1
            else:
                self.stats['requirements_from_prose'] += 1

        # ====================================================================
        # CHECK FOR MULTIPLE REQUIREMENTS IN ONE PARAGRAPH
        # ====================================================================
        # Some documents cram multiple requirements into one paragraph,
        # separated by numbers like "1. ... 2. ... 3. ..."
        elif self._has_multiple_requirements(text):
            multi_reqs = self._split_multiple_requirements(text, para_number)
            requirements.extend(multi_reqs)

        return requirements

    def _is_list_item(self, paragraph: Paragraph) -> bool:
        """
        PURPOSE:
            Check if a paragraph is a list item (numbered or bulleted).

        R EQUIVALENT:
            In officer, list items have specific style markers.

        PARAMETERS:
            paragraph (Paragraph): The paragraph to check

        RETURNS:
            bool: True if this is a list item

        WHY THIS APPROACH:
            Word stores list items as regular paragraphs with special
            numbering properties. We check the XML for list markers.
        """
        # Check for list numbering in the XML
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                return True

        # Also check style name for list styles
        style_name = paragraph.style.name if paragraph.style else ''
        list_styles = ['list', 'bullet', 'number']
        if any(ls in style_name.lower() for ls in list_styles):
            return True

        # Check text pattern (starts with bullet or number)
        text = paragraph.text.strip()
        if re.match(r'^[\u2022\u2023\u25E6\u2043\u2219\-\*]\s', text):
            return True  # Bullet characters
        if re.match(r'^\d+[.)]\s', text):
            return True  # Numbered: "1. " or "1) "
        if re.match(r'^[a-zA-Z][.)]\s', text):
            return True  # Lettered: "a. " or "a) "

        return False

    def _has_multiple_requirements(self, text: str) -> bool:
        """
        PURPOSE:
            Check if a paragraph contains multiple numbered requirements.

        PARAMETERS:
            text (str): The paragraph text

        RETURNS:
            bool: True if there appear to be multiple requirements

        WHY THIS APPROACH:
            Some authors write "1. First requirement 2. Second requirement"
            all in one paragraph. We detect this pattern.
        """
        # Look for multiple numbered items in the text
        # Pattern: digit + period/parenthesis + space, appearing 2+ times
        matches = re.findall(r'\d+[.)]\s+\S', text)
        return len(matches) >= 2

    def _split_multiple_requirements(self, text: str, base_para: int) -> list[dict]:
        """
        PURPOSE:
            Split a paragraph containing multiple numbered requirements.

        PARAMETERS:
            text (str): The paragraph text with multiple requirements
            base_para (int): Base paragraph number

        RETURNS:
            list[dict]: List of individual requirements

        WHY THIS APPROACH:
            We split on the number pattern and process each piece
            as a separate requirement.
        """
        requirements = []

        # Split on numbered patterns, keeping the delimiter
        parts = re.split(r'(?=\d+[.)]\s+)', text)
        parts = [p.strip() for p in parts if p.strip()]

        for i, part in enumerate(parts):
            # Skip if too short
            if len(part) < 10:
                continue

            # Check if it looks like a requirement
            part_lower = part.lower()
            has_indicator = any(ind in part_lower for ind in self.requirement_indicators)

            if has_indicator or re.match(r'^\d+[.)]\s', part):
                # Extract ID if present
                explicit_id = self._extract_requirement_id(part)

                requirement = {
                    'raw_text': part,
                    'description': self._clean_requirement_text(part),
                    'row_number': base_para,
                    'source_cell': f"Paragraph {base_para}, Item {i+1}",
                    'requirement_id': explicit_id or self._generate_id(base_para, i+1),
                    'priority': self._infer_priority(part),
                    'category': self._current_section or 'General',
                    'notes': '',
                    'title': '',
                    'status': ''
                }
                requirements.append(requirement)
                self.stats['requirements_from_prose'] += 1

        return requirements

    # ========================================================================
    # TABLE PARSING
    # ========================================================================

    def _parse_table(self, table: Table, table_number: int) -> list[dict]:
        """
        PURPOSE:
            Extract requirements from a Word table.
            Identifies header row and maps columns to requirement fields.

        R EQUIVALENT:
            Similar to reading a table with officer::docx_tables() and
            then cleaning it up with janitor::row_to_names().

        PARAMETERS:
            table (Table): The Word table object
            table_number (int): Table number for reference

        RETURNS:
            list[dict]: List of requirement dictionaries

        WHY THIS APPROACH:
            Tables are the most structured format for requirements.
            We identify the header row, map columns, and extract data.
        """
        requirements = []

        # Get all rows from the table
        rows = table.rows
        if len(rows) < 2:
            # Need at least header + one data row
            return requirements

        # ====================================================================
        # IDENTIFY HEADER ROW
        # ====================================================================
        header_row_idx, column_map = self._identify_table_structure(table)

        if column_map is None:
            # Couldn't identify structure - try to extract any requirements
            return self._parse_unstructured_table(table, table_number)

        # ====================================================================
        # EXTRACT DATA ROWS
        # ====================================================================
        for row_idx, row in enumerate(rows):
            # Skip header row
            if row_idx <= header_row_idx:
                continue

            # Extract cell values
            cells = row.cells
            cell_texts = [self._get_cell_text(cell) for cell in cells]

            # Skip empty rows
            if not any(cell_texts):
                continue

            # Map to requirement fields
            requirement = self._map_row_to_requirement(
                cell_texts, column_map, table_number, row_idx
            )

            if requirement:
                requirements.append(requirement)
                self.stats['requirements_from_tables'] += 1

        return requirements

    def _identify_table_structure(self, table: Table) -> tuple[int, Optional[dict]]:
        """
        PURPOSE:
            Identify the header row and map columns to requirement fields.

        PARAMETERS:
            table (Table): The Word table

        RETURNS:
            tuple: (header_row_index, column_mapping_dict or None)
                   column_mapping is like {'description': 2, 'priority': 3}

        WHY THIS APPROACH:
            Tables may have the header in row 0, 1, or even later.
            We scan rows looking for one that contains recognizable headers.
        """
        # Check first few rows for header patterns
        for row_idx in range(min(3, len(table.rows))):
            row = table.rows[row_idx]
            cells = row.cells
            cell_texts = [self._get_cell_text(cell).lower() for cell in cells]

            # Try to map columns
            column_map = {}
            for field, keywords in self.table_column_mappings.items():
                for col_idx, cell_text in enumerate(cell_texts):
                    if any(kw in cell_text for kw in keywords):
                        column_map[field] = col_idx
                        break

            # If we found at least a description column, this is our header
            if 'description' in column_map:
                return row_idx, column_map

            # Also accept if we found ID column (ID column implies it's a req table)
            if 'id' in column_map and len(column_map) >= 2:
                return row_idx, column_map

        # No clear header found
        return 0, None

    def _parse_unstructured_table(self, table: Table, table_number: int) -> list[dict]:
        """
        PURPOSE:
            Try to extract requirements from a table without clear structure.
            Falls back to treating each row as potential requirement text.

        PARAMETERS:
            table (Table): The Word table
            table_number (int): Table number for reference

        RETURNS:
            list[dict]: List of requirement dictionaries

        WHY THIS APPROACH:
            Some tables don't have headers or use non-standard column names.
            We still try to extract useful content by checking each cell
            for requirement indicators.
        """
        requirements = []

        for row_idx, row in enumerate(table.rows):
            # Concatenate all cell text in the row
            row_text = ' '.join(self._get_cell_text(cell) for cell in row.cells)
            row_text = row_text.strip()

            # Skip short/empty rows
            if len(row_text) < 15:
                continue

            # Check if this looks like a requirement
            text_lower = row_text.lower()
            has_indicator = any(ind in text_lower for ind in self.requirement_indicators)
            explicit_id = self._extract_requirement_id(row_text)

            if has_indicator or explicit_id:
                requirement = {
                    'raw_text': row_text,
                    'description': self._clean_requirement_text(row_text),
                    'row_number': row_idx,
                    'source_cell': f"Table {table_number}, Row {row_idx + 1}",
                    'requirement_id': explicit_id or self._generate_id(table_number * 100 + row_idx),
                    'priority': self._infer_priority(row_text),
                    'category': self._current_section or 'General',
                    'notes': f"From unstructured table {table_number}",
                    'title': '',
                    'status': ''
                }
                requirements.append(requirement)
                self.stats['requirements_from_tables'] += 1

        return requirements

    def _get_cell_text(self, cell: _Cell) -> str:
        """
        PURPOSE:
            Extract clean text from a table cell.
            Handles nested paragraphs and merged cells.

        PARAMETERS:
            cell (_Cell): The Word table cell

        RETURNS:
            str: Cleaned cell text

        WHY THIS APPROACH:
            Word cells can contain multiple paragraphs.
            We concatenate them with newlines preserved.
        """
        # Get text from all paragraphs in the cell
        paragraphs = cell.paragraphs
        texts = [p.text.strip() for p in paragraphs if p.text.strip()]
        return '\n'.join(texts)

    def _map_row_to_requirement(self, cell_texts: list[str], column_map: dict,
                                 table_number: int, row_idx: int) -> Optional[dict]:
        """
        PURPOSE:
            Convert a table row to a requirement dictionary using column mapping.

        PARAMETERS:
            cell_texts (list[str]): Text from each cell in the row
            column_map (dict): Mapping of field names to column indices
            table_number (int): Table number for reference
            row_idx (int): Row index in the table

        RETURNS:
            dict or None: Requirement dictionary, or None if row is invalid

        WHY THIS APPROACH:
            We use the column mapping to extract each field from the
            appropriate cell position.
        """
        # Get description (required)
        desc_col = column_map.get('description', 0)
        if desc_col >= len(cell_texts):
            return None

        description = cell_texts[desc_col].strip()
        if not description or len(description) < 5:
            return None

        # Get ID (optional - generate if missing)
        req_id = ''
        if 'id' in column_map and column_map['id'] < len(cell_texts):
            req_id = cell_texts[column_map['id']].strip()

        # If no explicit ID, try to extract from description
        if not req_id:
            req_id = self._extract_requirement_id(description)
        if not req_id:
            req_id = self._generate_id(table_number * 100 + row_idx)

        # Get priority (optional)
        priority = 'Medium'
        if 'priority' in column_map and column_map['priority'] < len(cell_texts):
            priority_text = cell_texts[column_map['priority']].strip()
            priority = self._normalize_priority(priority_text)

        # If no explicit priority, infer from description
        if priority == 'Medium':
            inferred = self._infer_priority(description)
            if inferred != 'Medium':
                priority = inferred

        # Get category (optional - prefer section context)
        category = self._current_section or 'General'
        if 'category' in column_map and column_map['category'] < len(cell_texts):
            cat_text = cell_texts[column_map['category']].strip()
            if cat_text:
                category = cat_text

        # Get notes (optional)
        notes = ''
        if 'notes' in column_map and column_map['notes'] < len(cell_texts):
            notes = cell_texts[column_map['notes']].strip()

        # Get status (optional)
        status = ''
        if 'status' in column_map and column_map['status'] < len(cell_texts):
            status = cell_texts[column_map['status']].strip()

        return {
            'raw_text': description,
            'description': self._clean_requirement_text(description),
            'row_number': row_idx,
            'source_cell': f"Table {table_number}, Row {row_idx + 1}",
            'requirement_id': req_id,
            'priority': priority,
            'category': category,
            'notes': notes,
            'title': '',
            'status': status
        }

    # ========================================================================
    # COMMENTS EXTRACTION
    # ========================================================================

    def _extract_comments(self) -> list[dict]:
        """
        PURPOSE:
            Extract comments from the Word document.
            Comments often contain clarifications or additional requirements.

        R EQUIVALENT:
            officer doesn't have direct comment support; this would
            require parsing the XML directly.

        RETURNS:
            list[dict]: Requirements extracted from comments

        WHY THIS APPROACH:
            Word comments are stored in a separate XML part.
            We access them through the document's XML structure.
        """
        requirements = []

        try:
            # Access the comments part of the document
            # WHY: Comments are stored separately in the .docx package
            comments_part = self.document.part.package.part_related_by(
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
            )

            if comments_part is None:
                return requirements

            # Parse the comments XML
            from lxml import etree
            comments_xml = etree.fromstring(comments_part.blob)

            # Namespace for Word comments
            w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

            for comment in comments_xml.findall(f'.//{w_ns}comment'):
                # Get comment text
                comment_texts = comment.findall(f'.//{w_ns}t')
                comment_text = ' '.join(t.text for t in comment_texts if t.text)
                comment_text = comment_text.strip()

                if not comment_text:
                    continue

                # Check if this comment looks like a requirement
                text_lower = comment_text.lower()
                has_indicator = any(ind in text_lower for ind in self.requirement_indicators)

                if has_indicator:
                    comment_id = comment.get(f'{w_ns}id', 'unknown')
                    author = comment.get(f'{w_ns}author', 'unknown')

                    requirement = {
                        'raw_text': comment_text,
                        'description': self._clean_requirement_text(comment_text),
                        'row_number': 0,
                        'source_cell': f"Comment by {author}",
                        'requirement_id': f"CMT-{comment_id}",
                        'priority': self._infer_priority(comment_text),
                        'category': 'From Comments',
                        'notes': f"Extracted from document comment by {author}",
                        'title': '',
                        'status': ''
                    }
                    requirements.append(requirement)
                    self.stats['comments_extracted'] += 1

        except Exception:
            # Comments extraction is optional - don't fail the whole parse
            pass

        return requirements

    # ========================================================================
    # TEXT PROCESSING UTILITIES
    # ========================================================================

    def _extract_requirement_id(self, text: str) -> str:
        """
        PURPOSE:
            Extract an explicit requirement ID from text.

        PARAMETERS:
            text (str): Text that might start with an ID like "REQ-001"

        RETURNS:
            str: Extracted ID or empty string if not found

        WHY THIS APPROACH:
            We check against known ID patterns at the start of text.
            IDs typically appear at the beginning.
        """
        text = text.strip()

        for pattern in self.id_patterns:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).upper()

        return ''

    def _generate_id(self, number: int, sub_number: int = 0) -> str:
        """
        PURPOSE:
            Generate a requirement ID when none is present.

        PARAMETERS:
            number (int): Primary sequence number
            sub_number (int): Sub-number for items within a paragraph

        RETURNS:
            str: Generated ID like "WD-001" or "WD-001-A"

        WHY THIS APPROACH:
            We use "WD" prefix to indicate Word-derived requirements.
            This distinguishes them from IDs in other sources.
        """
        if sub_number > 0:
            # Convert sub_number to letter (1=A, 2=B, etc.)
            suffix = chr(64 + sub_number)  # 65 is 'A'
            return f"WD-{number:03d}-{suffix}"
        return f"WD-{number:03d}"

    def _is_non_requirement_text(self, text: str) -> bool:
        """
        PURPOSE:
            Check if text is likely NOT a requirement (overview, intro, etc.)

        PARAMETERS:
            text (str): The text to check

        RETURNS:
            bool: True if this should be skipped

        WHY THIS APPROACH:
            Documents often have introductory text that contains requirement
            keywords but isn't actually a requirement. We filter these out
            to avoid noise.
        """
        text_lower = text.lower().strip()

        # Patterns that indicate non-requirement text
        non_req_patterns = [
            r'^this\s+document\s+(outlines|describes|defines|specifies)',
            r'^the\s+following\s+(requirements|section|document)',
            r'^this\s+section\s+(describes|contains|covers)',
            r'^see\s+(appendix|section|table)',
            r'^refer\s+to\s+',
            r'^note:\s*',
            r'^for\s+more\s+information',
            r'^the\s+purpose\s+of\s+this',
            r'^overview\s*:',
            r'^introduction\s*:',
            r'^scope\s*:',
            r'^background\s*:',
            r'^definitions\s*:',
            r'^glossary\s*:',
            r'^revision\s+history',
            r'^document\s+version',
            r'^table\s+of\s+contents',
            r'^toc\s*$',
            r'^page\s+\d+',
            r'^\d+\s+of\s+\d+$',  # "3 of 10"
        ]

        for pattern in non_req_patterns:
            if re.match(pattern, text_lower):
                return True

        # Very short text without requirement indicator (likely a heading)
        if len(text) < 30 and not any(ind in text_lower for ind in self.requirement_indicators):
            # Check if it's just a title-like phrase (no verb)
            words = text_lower.split()
            if len(words) <= 4:
                return True

        return False

    def _clean_requirement_text(self, text: str) -> str:
        """
        PURPOSE:
            Clean up requirement text by removing noise.

        PARAMETERS:
            text (str): Raw requirement text

        RETURNS:
            str: Cleaned text

        WHY THIS APPROACH:
            We remove list markers, extra whitespace, and normalize
            the text for consistent output.
        """
        cleaned = text.strip()

        # Remove leading list markers (bullets, numbers)
        cleaned = re.sub(r'^[\u2022\u2023\u25E6\u2043\u2219\-\*]\s*', '', cleaned)
        cleaned = re.sub(r'^\d+[.)]\s*', '', cleaned)
        cleaned = re.sub(r'^[a-zA-Z][.)]\s*', '', cleaned)

        # Remove leading requirement IDs (already captured separately)
        for pattern in self.id_patterns:
            cleaned = re.sub(f'^{pattern}[:\\s]+', '', cleaned, flags=re.IGNORECASE)

        # Collapse multiple spaces/newlines
        cleaned = re.sub(r'\s+', ' ', cleaned)

        # Remove trailing punctuation if incomplete sentence
        cleaned = cleaned.strip()

        return cleaned

    def _infer_priority(self, text: str) -> str:
        """
        PURPOSE:
            Infer priority from keywords in the text.

        PARAMETERS:
            text (str): Text to analyze

        RETURNS:
            str: Priority level (Critical, High, Medium, Low)

        WHY THIS APPROACH:
            Same logic as other parsers for consistency.
            We look for priority keywords in the text.
        """
        text_lower = text.lower()

        for priority, keywords in self.priority_keywords.items():
            if any(kw in text_lower for kw in keywords):
                return priority.capitalize()

        return 'Medium'

    def _normalize_priority(self, priority_text: str) -> str:
        """
        PURPOSE:
            Normalize explicit priority text to standard values.

        PARAMETERS:
            priority_text (str): Raw priority text from table cell

        RETURNS:
            str: Normalized priority (Critical, High, Medium, Low)

        WHY THIS APPROACH:
            Tables may use various priority representations.
            We normalize them to our standard set.
        """
        text_lower = priority_text.lower().strip()

        # Direct matches
        if text_lower in ['critical', 'p0', 'p1']:
            return 'Critical'
        if text_lower in ['high', 'p2', '1', 'must']:
            return 'High'
        if text_lower in ['medium', 'p3', '2', 'should', 'normal']:
            return 'Medium'
        if text_lower in ['low', 'p4', '3', 'could', 'nice to have']:
            return 'Low'

        # Check for keywords
        return self._infer_priority(priority_text)

    # ========================================================================
    # DEDUPLICATION
    # ========================================================================

    def _deduplicate_requirements(self, requirements: list[dict]) -> list[dict]:
        """
        PURPOSE:
            Remove duplicate requirements (same text appearing multiple times).

        PARAMETERS:
            requirements (list[dict]): All extracted requirements

        RETURNS:
            list[dict]: Deduplicated list

        WHY THIS APPROACH:
            Word documents sometimes repeat content (in summary tables,
            appendices, etc.). We keep the first occurrence.
        """
        seen_texts = set()
        unique_requirements = []

        for req in requirements:
            # Normalize text for comparison
            normalized = req['description'].lower().strip()
            normalized = re.sub(r'\s+', ' ', normalized)

            if normalized not in seen_texts and len(normalized) > 10:
                seen_texts.add(normalized)
                unique_requirements.append(req)

        return unique_requirements

    # ========================================================================
    # PUBLIC UTILITIES
    # ========================================================================

    def get_stats(self) -> dict:
        """
        PURPOSE:
            Return parsing statistics.

        RETURNS:
            dict: Statistics about what was parsed

        WHY THIS APPROACH:
            Helps users understand what was extracted and spot issues.
        """
        return self.stats.copy()

    def get_sections(self) -> list[str]:
        """
        PURPOSE:
            Return list of sections found in the document.

        RETURNS:
            list[str]: Section headings in order

        WHY THIS APPROACH:
            Useful for understanding document structure.
        """
        # Parse if not already done
        if self.document is None:
            self.parse()

        sections = []
        for para in self.document.paragraphs:
            if self._is_heading(para) and para.text.strip():
                sections.append(para.text.strip())

        return sections


# ============================================================================
# STANDALONE TEST
# ============================================================================

if __name__ == "__main__":
    """
    PURPOSE:
        Test the WordParser independently.

    USAGE:
        python3 parsers/word_parser.py "path/to/document.docx"
        python3 parsers/word_parser.py  # Uses sample data

    WHY THIS APPROACH:
        Allows testing the parser without running the full pipeline.
        Creates a simple test document if none provided.
    """
    import sys

    print("=" * 70)
    print("WORD PARSER TEST")
    print("=" * 70)

    # Check for command line argument
    if len(sys.argv) > 1:
        test_file = sys.argv[1]
        print(f"\nTesting with: {test_file}")

        try:
            parser = WordParser(test_file)
            requirements = parser.parse()

            print(f"\nParsed {len(requirements)} requirements")
            print(f"\nStatistics:")
            for key, value in parser.get_stats().items():
                print(f"  {key}: {value}")

            # Show sections found
            sections = parser.get_sections()
            if sections:
                print(f"\nSections found: {len(sections)}")
                for i, section in enumerate(sections[:5]):
                    print(f"  - {section}")
                if len(sections) > 5:
                    print(f"  ... and {len(sections) - 5} more")

            # Show sample requirements
            print(f"\nSample requirements:")
            for req in requirements[:5]:
                desc = req['description'][:60] + '...' if len(req['description']) > 60 else req['description']
                print(f"  {req['requirement_id']}: {desc}")
                print(f"      Category: {req['category']} | Priority: {req['priority']}")
                print(f"      Source: {req['source_cell']}")

        except ImportError as e:
            print(f"\nError: {e}")
            print("Install python-docx with: pip install python-docx")
        except Exception as e:
            print(f"\nError parsing document: {e}")
            raise

    else:
        # No file provided - create a simple test
        print("\nNo file provided. Creating test document...")

        if not DOCX_AVAILABLE:
            print("Error: python-docx not installed.")
            print("Install with: pip install python-docx")
            sys.exit(1)

        # Create a test document
        test_path = Path("inputs/word/test_requirements.docx")
        test_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Add title
        doc.add_heading("Sample Requirements Document", 0)

        # Add overview section
        doc.add_heading("1. Overview", level=1)
        doc.add_paragraph(
            "This document outlines the requirements for the new user management system."
        )

        # Add functional requirements section with table
        doc.add_heading("2. Functional Requirements", level=1)
        doc.add_heading("2.1 Authentication", level=2)

        # Add a requirements table
        table = doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'

        # Header row
        headers = ['ID', 'Description', 'Priority', 'Notes']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        # Data rows
        data = [
            ['REQ-001', 'The system shall allow users to login with email and password', 'High', 'Core feature'],
            ['REQ-002', 'Users must be able to reset their password via email', 'High', 'Security requirement'],
            ['REQ-003', 'The system should support two-factor authentication', 'Medium', 'Future enhancement'],
        ]
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_data in enumerate(row_data):
                table.rows[row_idx + 1].cells[col_idx].text = cell_data

        # Add prose requirements
        doc.add_heading("2.2 User Management", level=2)
        doc.add_paragraph(
            "The system shall allow administrators to create new user accounts. "
            "Administrators must be able to assign roles to users. "
            "The system should log all user management activities."
        )

        # Add list requirements
        doc.add_heading("3. Non-Functional Requirements", level=1)
        doc.add_paragraph("The following requirements apply to system performance:")
        doc.add_paragraph("1. The system shall respond to user actions within 2 seconds", style='List Number')
        doc.add_paragraph("2. The system must handle at least 1000 concurrent users", style='List Number')
        doc.add_paragraph("3. System availability should be 99.9% uptime", style='List Number')

        # Save the document
        doc.save(str(test_path))
        print(f"Created test document: {test_path}")

        # Now parse it
        print(f"\nParsing test document...")
        parser = WordParser(str(test_path))
        requirements = parser.parse()

        print(f"\nParsed {len(requirements)} requirements")
        print(f"\nStatistics:")
        for key, value in parser.get_stats().items():
            print(f"  {key}: {value}")

        print(f"\nAll requirements:")
        for req in requirements:
            desc = req['description'][:60] + '...' if len(req['description']) > 60 else req['description']
            print(f"  {req['requirement_id']}: {desc}")
            print(f"      Category: {req['category']} | Priority: {req['priority']}")
            print(f"      Source: {req['source_cell']}")
            print()
